import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color="000000")
    set_paragraph_spacing(paragraph)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10, italic=True, color="666666")
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.15)
    return paragraph


def add_prompt_block(doc, label, text):
    heading = doc.add_paragraph()
    run = heading.add_run(label)
    set_run_font(run, size=10.5, bold=True, color="000000")
    set_paragraph_spacing(heading, before=8, after=4, line=1.15)

    for chunk in [part.strip() for part in text.split("\n\n") if part.strip()]:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(chunk)
        set_run_font(run, size=9.5, color="333333")
        set_paragraph_spacing(paragraph, before=0, after=4, line=1.08)
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.right_indent = Inches(0.12)
        shade_paragraph(paragraph, "F3F3F3")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [
        ("Heading 1", 20, 20, 6),
        ("Heading 2", 16, 18, 6),
        ("Heading 3", 14, 16, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def build_doc(data, output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)

    title = doc.add_paragraph()
    title_run = title.add_run(data["title"])
    set_run_font(title_run, size=26, bold=False, color="000000")
    set_paragraph_spacing(title, before=0, after=3, line=1.15)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f'{data["tag"]} · {data["date"]}')
    set_run_font(meta_run, size=10.5, color="555555")
    set_paragraph_spacing(meta, before=0, after=14, line=1.15)

    if data.get("asideText"):
        doc.add_heading("Краткое описание", level=2)
        add_paragraph(doc, data["asideText"])

    for section_data in data["sections"]:
        title_text = section_data.get("title", "").strip()
        if title_text:
            doc.add_heading(title_text, level=1)

        for item in section_data["paragraphs"]:
            if item.get("text"):
                add_paragraph(doc, item["text"])

            for caption in item.get("captionsAfter", []):
                add_caption(doc, caption)

            for prompt in item.get("promptsAfter", []):
                add_prompt_block(doc, prompt.get("label", "Prompt"), prompt.get("text", ""))

    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_part2_docx.py article.json output.docx")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    data = json.loads(input_path.read_text())
    build_doc(data, output_path)


if __name__ == "__main__":
    main()
