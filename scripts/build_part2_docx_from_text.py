import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


EDITORIAL_MARKERS = (
    "тут ",
    "убери",
    "замени",
    "переп",
    "сдела",
    "добав",
    "помен",
    "ссылк",
    "копир",
    "обезлич",
    "расшиф",
    "подписи",
    "фото",
    "интерактив",
    "при нажатии",
    "плашк",
    "встав",
)

KNOWN_HEADINGS = {
    "Краткое описание",
    "Что разбираем",
    "Смена ракурса",
    "Среда выполнения и генеративные модели",
    "Эксперимент 1: модель в студии",
    "Товарная сцена с большим количеством деталей",
    "Reference description",
    "Prompt",
    "Вывод",
}


def set_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_font(run, size=11, bold=None, italic=None, color="000000", name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(paragraph, fill="F3F3F3"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def has_editorial_note(text):
    lowered = text.lower()
    return any(marker in lowered for marker in EDITORIAL_MARKERS)


def remove_editorial_parentheses(text):
    def repl(match):
        inner = match.group(1).strip()
        lowered = inner.lower()
        if "информация актуальна" in lowered:
            return f" ({inner})"
        if has_editorial_note(inner) or inner.startswith("/Users/"):
            return ""
        return f" ({inner})"

    # Repeat to handle multiple notes in one paragraph.
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\(([^()]*)\)", repl, text)
    return text


def clean_line(line):
    line = line.strip()
    if not line:
        return ""

    # Skip purely editorial/layout instructions.
    low = line.lower()
    if low.startswith(("тут давай", "куда вставить", "image -", "gemini 3 -", "chatgpt")) and has_editorial_note(line):
        return ""
    if "убери этот абзац" in low:
        return ""

    line = line.replace("ChatGPT Imagen 2.0 Edit", "ChatGPT Imagen 2.0")
    line = line.replace("Claude Code", "Claude")
    line = line.replace("конкретный технический референс", "конкретный референс")
    line = line.replace("Для финальной ручной доводки  *все равно* ", "Для финальной ручной доводки ")
    line = line.replace("Для финальной ручной доводки *все равно* ", "Для финальной ручной доводки ")
    line = line.replace("Консультация или генеративная работа под задачу.", "Консультация или работа под задачу.")
    line = re.sub(r"\s*Беру проекты разной сложности\.\s*", " ", line)

    # Remove inline editor notes and markdown-ish invisible separators.
    line = remove_editorial_parentheses(line)
    line = re.sub(r"\s+-\s+убери.*$", "", line, flags=re.IGNORECASE)
    line = re.sub(r"\s+-\s+замени.*$", "", line, flags=re.IGNORECASE)
    line = line.replace("\u2028", " ").replace("\u2029", " ")
    line = line.replace(" ", " ")
    line = re.sub(r"\s{2,}", " ", line).strip()
    return line


def looks_like_prompt(text):
    if text in {"Prompt", "Reference description"}:
        return False
    cyr = sum(1 for ch in text if "\u0400" <= ch <= "\u04ff")
    latin = sum(1 for ch in text if ch.isascii() and ch.isalpha())
    starters = (
        "Ultra-realistic",
        "Use the",
        "The girl",
        "The room",
        "Negative prompt",
        "SECOND IMAGE",
        "Create ",
        "Only ",
        "Visual style",
        "Lighting",
        "Camera",
        "The main objects",
        "There is",
        "A red-and-white",
        "The background",
        "The shadows",
        "The dominant color",
        "Both cartons",
        "TARGET CARTON",
        "CARTON SHAPE",
        "LIGHTING INTEGRATION",
        "MATERIAL INTEGRATION",
        "PERSPECTIVE",
        "BACKGROUND AND SURROUNDINGS",
        "FINAL RESULT",
    )
    return text.startswith(starters) or (latin > 50 and latin > cyr * 3)


def looks_like_heading(text, index):
    if text in KNOWN_HEADINGS:
        return True
    if index < 2:
        return False
    if looks_like_prompt(text):
        return False
    if len(text) > 90:
        return False
    if text.endswith((".", ":", ";", ",")):
        return False
    if re.search(r"[a-zA-Z]{15,}", text):
        return False
    return True


def add_rich_text(paragraph, text, size=11, color="000000"):
    # Minimal markdown emphasis support for *italic* fragments.
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        italic = part.startswith("*") and part.endswith("*") and len(part) > 2
        value = part[1:-1] if italic else part
        run = paragraph.add_run(value)
        set_font(run, size=size, italic=italic, color=color)


def add_body(doc, text):
    p = doc.add_paragraph()
    add_rich_text(p, text, size=11)
    set_spacing(p)


def add_caption(doc, text):
    p = doc.add_paragraph()
    add_rich_text(p, text, size=10, color="666666")
    set_spacing(p, before=0, after=6)


def add_prompt(doc, text):
    p = doc.add_paragraph()
    add_rich_text(p, text, size=9.5, color="333333")
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    set_spacing(p, before=0, after=4, line=1.08)
    shade(p)


def build(input_path, output_path):
    raw_lines = input_path.read_text(encoding="utf-8").splitlines()
    lines = [clean_line(line) for line in raw_lines]
    lines = [line for line in lines if line]

    doc = Document()
    configure(doc)

    title = lines[0] if lines else "Контролирование пространства и планировки. Часть 2"
    meta = lines[1] if len(lines) > 1 else "Туториал · 8 июня 2026"

    p = doc.add_paragraph()
    run = p.add_run(title)
    set_font(run, size=26, color="000000")
    set_spacing(p, after=3)

    p = doc.add_paragraph()
    run = p.add_run(meta)
    set_font(run, size=10.5, color="555555")
    set_spacing(p, after=14)

    previous_was_prompt_label = False
    for index, line in enumerate(lines[2:], start=2):
        if looks_like_heading(line, index):
            level = 2 if line in {"Prompt", "Reference description"} else 1
            doc.add_heading(line, level=level)
            previous_was_prompt_label = line in {"Prompt", "Reference description"}
            continue

        if looks_like_prompt(line) or previous_was_prompt_label:
            add_prompt(doc, line)
            previous_was_prompt_label = looks_like_prompt(line)
        elif line.startswith("Подписи колонок:") or line.startswith("Этап "):
            add_caption(doc, line)
            previous_was_prompt_label = False
        else:
            add_body(doc, line)
            previous_was_prompt_label = False

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_part2_docx_from_text.py source.txt output.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()
